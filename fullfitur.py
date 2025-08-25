import os
import re
import sys
import logging
import json
import requests
import tempfile
import shutil
from datetime import datetime
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, # Ditambahkan untuk post_init
    ApplicationBuilder, CommandHandler, MessageHandler, filters,
    ContextTypes, ConversationHandler, CallbackQueryHandler
)
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from PIL import Image, ImageDraw, ImageFont

# Muat variabel dari file .env
load_dotenv()

# Konfigurasi Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('rembes_bot.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class Config:
    TOKEN = os.getenv('BOT_TOKEN')
    GROUP_CHAT_ID = os.getenv('GROUP_CHAT_ID')
    BUNNY_STORAGE_ZONE_NAME = os.getenv('BUNNY_STORAGE_ZONE_NAME')
    BUNNY_ACCESS_KEY = os.getenv('BUNNY_ACCESS_KEY')
    BUNNY_REGION = os.getenv('BUNNY_REGION', '')
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    CONFIG_FILE = os.path.join(BASE_DIR, 'config.json')

class BunnyStorage:
    def __init__(self, zone_name, access_key, region=''):
        base_url = "storage.bunnycdn.com"
        if region and region.lower() != 'de':
            base_url = f"{region}.{base_url}"
        self.api_url = f"https://{base_url}/{zone_name}/"
        self.headers = {"AccessKey": access_key}

    def check_connection(self):
        return self.list_files('') is not None

    def list_files(self, remote_path):
        try:
            response = requests.get(self.api_url + remote_path, headers=self.headers)
            if response.status_code == 200: return response.json()
            elif response.status_code == 404: return []
            logger.error(f"Bunny API Error (List): {response.status_code} - {response.text}"); return None
        except requests.exceptions.RequestException as e:
            logger.error(f"Bunny Connection Error (List): {e}"); return None

    def upload_file(self, local_file_path, remote_file_path):
        headers = self.headers.copy(); headers["Content-Type"] = "application/octet-stream"
        try:
            with open(local_file_path, 'rb') as f:
                response = requests.put(self.api_url + remote_file_path, headers=headers, data=f)
                return response.status_code == 201
        except requests.exceptions.RequestException as e:
            logger.error(f"Bunny Connection Error (Upload): {e}"); return False

    def download_file(self, remote_file_path, local_file_path):
        try:
            response = requests.get(self.api_url + remote_file_path, headers=self.headers, stream=True)
            if response.status_code == 200:
                with open(local_file_path, 'wb') as f: shutil.copyfileobj(response.raw, f)
                return True
            return False
        except requests.exceptions.RequestException as e:
            logger.error(f"Bunny Connection Error (Download): {e}"); return False

    def delete_file(self, remote_file_path):
        try:
            response = requests.delete(self.api_url + remote_file_path, headers=self.headers)
            if response.status_code == 200:
                logger.info(f"File '{remote_file_path}' berhasil dihapus."); return True, "✅ File berhasil dihapus."
            elif response.status_code == 404:
                logger.warning(f"Gagal hapus, file tidak ditemukan: '{remote_file_path}'."); return True, "ℹ️ File tidak ditemukan (mungkin sudah dihapus)."
            logger.error(f"Bunny API Error (Delete): {response.status_code} - {response.text}"); return False, f"❌ Gagal menghapus file. Error: {response.status_code}"
        except requests.exceptions.RequestException as e:
            logger.error(f"Bunny Connection Error (Delete): {e}"); return False, "❌ Gagal terhubung ke storage."

class RembesBot:
    def __init__(self):
        self.config = Config()
        self.storage = BunnyStorage(self.config.BUNNY_STORAGE_ZONE_NAME, self.config.BUNNY_ACCESS_KEY, self.config.BUNNY_REGION)
        self.commands = self._load_commands()
        self.application = None
        self.start_time = datetime.now()
        (
            self.GET_PHOTO, self.GET_KETERANGAN, self.GET_BIAYA, self.ASK_CONTINUE,
            self.CHOOSE_DELETE_CATEGORY, self.CHOOSE_DELETE_FILE,
            self.CHOOSE_EDIT_CATEGORY, self.CHOOSE_EDIT_FILE, self.CHOOSE_EDIT_FIELD, self.GET_NEW_VALUE
        ) = range(10)

    # --- Fungsi Utilitas ---
    def _load_commands(self):
        try:
            with open(self.config.CONFIG_FILE, 'r') as f: return json.load(f).get('commands', [])
        except (FileNotFoundError, json.JSONDecodeError):
            default_commands = ["grab", "mrt", "bensin", "parkir", "lembur"]; self._save_commands(default_commands)
            return default_commands

    def _save_commands(self, commands_list):
        with open(self.config.CONFIG_FILE, 'w') as f: json.dump({'commands': sorted(list(set(commands_list)))}, f, indent=2)

    def _get_current_period(self):
        now = datetime.now(); month = now.month - 1 if now.day < 25 else now.month
        year = now.year if now.month > 1 or now.day >= 25 else now.year - 1
        if month == 0: month = 12
        return f"{year}-{month:02d}", year, month

    def _format_uptime(self, duration):
        days, rem = divmod(duration.total_seconds(), 86400); hours, rem = divmod(rem, 3600); minutes, _ = divmod(rem, 60)
        return f"{int(days)} hari, {int(hours)} jam, {int(minutes)} menit"

    def _create_placeholder_image(self, text, path):
        img = Image.new('RGB', (400, 100), color=(80, 80, 80)); d = ImageDraw.Draw(img)
        try: font = ImageFont.truetype("arial.ttf", 15)
        except IOError: font = ImageFont.load_default()
        d.text((10,10), f"Tidak Ada Bukti Gambar\n\n{text}", fill=(255,255,255), font=font)
        img.save(path)

    # --- Handler Perintah Standar ---
    async def start_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        self.commands = self._load_commands()
        command_list_str = "".join([f"`{cmd}` " for cmd in self.commands])
        welcome_message = (
            "🤖 *Bot Rembesan v3.1!*\n\n"
            f"Kategori tersedia: {command_list_str}\n\n"
            "*Perintah Utama:*\n"
            "`/status` - Cek status bot\n"
            "`/summary` - Lihat ringkasan total biaya\n"
            "`/list` - Lihat rincian data\n"
            "`/export` - Ekspor data ke Word\n"
            "`/hapus` - Hapus data\n"
            "`/edit` - Edit data\n\n"
            "*Manajemen:*\n"
            "`/tambah_kategori` | `/hapus_kategori`\n"
            "`/batal` - Membatalkan proses"
        )
        await update.message.reply_text(welcome_message, parse_mode='Markdown')

    async def status_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        uptime = self._format_uptime(datetime.now() - self.start_time)
        bunny_status = "✅ Terhubung" if self.storage.check_connection() else "❌ Gagal terhubung"
        period, _, _ = self._get_current_period()
        status_message = (
            f"🤖 *Status Bot*\n\n"
            f"🟢 Status: *Online*\n"
            f"⏱️ Waktu Aktif: *{uptime}*\n"
            f"☁️ Koneksi Storage: *{bunny_status}*\n"
            f"🗂️ Jumlah Kategori: *{len(self.commands)}*\n"
            f"🗓️ Periode Aktif: *{period}*"
        )
        await update.message.reply_text(status_message, parse_mode='Markdown')
        
    async def summary_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        period, _, _ = self._get_current_period()
        summary = {}; grand_total = 0
        message = [f"📊 *Ringkasan Biaya Periode {period}*"]
        for cmd in self.commands:
            remote_path = f"{cmd}/{period}/"; files = self.storage.list_files(remote_path)
            if files:
                try:
                    total = sum(int(f['ObjectName'].replace(".jpg","").split("_")[-1]) for f in files)
                    summary[cmd] = total; grand_total += total
                except (ValueError, IndexError): continue
        if not summary:
            await update.message.reply_text("Tidak ada data untuk ditampilkan."); return
        for cmd, total in summary.items(): message.append(f"- {cmd.upper()}: `Rp {total:,}`")
        message.append(f"\n*Grand Total: `Rp {grand_total:,}`*")
        await update.message.reply_text("\n".join(message), parse_mode='Markdown')

    async def add_category_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        try:
            new_cmd = context.args[0].lower().strip()
            if not new_cmd.isalpha(): await update.message.reply_text("❌ Nama kategori hanya boleh berisi huruf."); return
            if new_cmd in self.commands: await update.message.reply_text(f"⚠️ Kategori '{new_cmd}' sudah ada."); return
            self.commands.append(new_cmd); self._save_commands(self.commands)
            await update.message.reply_text(f"✅ Kategori '{new_cmd}' ditambahkan!\n\n‼️ *PENTING: Mohon restart bot agar perintah baru bisa digunakan.*", parse_mode='Markdown')
        except IndexError: await update.message.reply_text("Format: `/tambah_kategori <nama_kategori>`")

    async def remove_category_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        try:
            cmd_to_remove = context.args[0].lower().strip()
            if cmd_to_remove not in self.commands: await update.message.reply_text(f"⚠️ Kategori '{cmd_to_remove}' tidak ditemukan."); return
            self.commands.remove(cmd_to_remove); self._save_commands(self.commands)
            await update.message.reply_text(f"✅ Kategori '{cmd_to_remove}' dihapus!\n\n‼️ *PENTING: Mohon restart bot untuk menerapkan perubahan.*", parse_mode='Markdown')
        except IndexError: await update.message.reply_text("Format: `/hapus_kategori <nama_kategori>`")

    # --- Alur Percakapan Tambah Data ---
    async def start_reimbursement_flow(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        command = update.message.text[1:].lower(); context.user_data['category'] = command
        if command == 'lembur':
            context.user_data['no_photo'] = True
            await update.message.reply_text("Baik, masukkan keterangan untuk *LEMBUR*.\n\nKetik /batal untuk membatalkan.", parse_mode='Markdown')
            return self.GET_KETERANGAN
        else:
            context.user_data['no_photo'] = False
            await update.message.reply_text(f"Baik, silakan kirim foto/bukti untuk *{command.upper()}*.\n\nKetik /batal.", parse_mode='Markdown')
            return self.GET_PHOTO

    async def get_photo(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        photo_file = await update.message.photo[-1].get_file(); context.user_data['photo_file_id'] = photo_file.file_id
        await update.message.reply_text("✅ Foto diterima. Sekarang, masukkan keterangan singkat."); return self.GET_KETERANGAN

    async def get_keterangan(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        context.user_data['keterangan'] = update.message.text
        await update.message.reply_text("✅ Keterangan disimpan. Terakhir, masukkan jumlah biaya (hanya angka)."); return self.GET_BIAYA

    async def get_biaya_and_save(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        biaya = update.message.text
        if not biaya.isdigit():
            await update.message.reply_text("❌ Biaya harus berupa angka. Silakan masukkan lagi."); return self.GET_BIAYA
        await update.message.reply_text("⏳ Sedang memproses dan mengunggah data...")
        try:
            category = context.user_data['category']; keterangan = context.user_data['keterangan']
            period, _, _ = self._get_current_period(); safe_keterangan = re.sub(r'[^\w\-_\.]', '_', keterangan)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S'); file_name = f"{timestamp}_{safe_keterangan}_{biaya}.jpg"
            remote_path = f"{category}/{period}/{file_name}"
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as temp_file:
                if context.user_data.get('no_photo'):
                    self._create_placeholder_image(f"{keterangan}\nRp {int(biaya):,}", temp_file.name)
                else:
                    photo_file = await context.bot.get_file(context.user_data['photo_file_id'])
                    await photo_file.download_to_drive(temp_file.name)
                if self.storage.upload_file(temp_file.name, remote_path):
                    if not context.user_data.get('no_photo'):
                        if self.config.GROUP_CHAT_ID: await context.bot.send_photo(self.config.GROUP_CHAT_ID, context.user_data['photo_file_id'], caption=f"☁️ Data Baru:\n`{category.upper()}` | `{safe_keterangan}` | `Rp {int(biaya):,}`", parse_mode='Markdown')
                    else:
                        if self.config.GROUP_CHAT_ID: await context.bot.send_message(self.config.GROUP_CHAT_ID, f"☁️ Data Lembur (Tanpa Foto):\n`{category.upper()}` | `{safe_keterangan}` | `Rp {int(biaya):,}`", parse_mode='Markdown')
                    buttons = [[InlineKeyboardButton("Ya", callback_data="continue_yes"), InlineKeyboardButton("Tidak", callback_data="continue_no")]]
                    await update.message.reply_text(f"✅ Data berhasil disimpan. Ingin menambahkan data lagi untuk kategori *{category.upper()}*?", reply_markup=InlineKeyboardMarkup(buttons), parse_mode='Markdown')
                    return self.ASK_CONTINUE
                else: await update.message.reply_text("❌ Gagal mengunggah file ke cloud."); context.user_data.clear(); return ConversationHandler.END
            os.unlink(temp_file.name)
        except Exception as e:
            logger.error(f"Kesalahan saat menyimpan data: {e}"); await update.message.reply_text("❌ Terjadi kesalahan internal saat menyimpan.")
            context.user_data.clear(); return ConversationHandler.END

    async def ask_continue(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query; await query.answer()
        category = context.user_data['category']
        context.user_data.pop('photo_file_id', None); context.user_data.pop('keterangan', None)
        if query.data == 'continue_yes':
            if context.user_data.get('no_photo'):
                await query.edit_message_text(f"Baik, masukkan keterangan lagi untuk *{category.upper()}*.", parse_mode='Markdown'); return self.GET_KETERANGAN
            else:
                await query.edit_message_text(f"Baik, silakan kirim foto/bukti lagi untuk *{category.upper()}*.", parse_mode='Markdown'); return self.GET_PHOTO
        else:
            await query.edit_message_text("Baik, proses selesai."); context.user_data.clear()
            return ConversationHandler.END

    # --- Alur Percakapan Hapus Data ---
    async def start_delete_flow(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        buttons = [[InlineKeyboardButton(cmd.upper(), callback_data=f"delcat_{cmd}")] for cmd in self.commands]
        await update.message.reply_text("Dari kategori mana Anda ingin menghapus data?", reply_markup=InlineKeyboardMarkup(buttons))
        return self.CHOOSE_DELETE_CATEGORY

    async def choose_delete_category(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query; await query.answer(); category = query.data.split('_')[1]
        context.user_data['delete_category'] = category; period, _, _ = self._get_current_period()
        remote_path = f"{category}/{period}/"; files = self.storage.list_files(remote_path)
        if not files:
            await query.edit_message_text(f"Tidak ada file untuk dihapus di kategori *{category.upper()}*.", parse_mode='Markdown')
            context.user_data.clear(); return ConversationHandler.END
        context.user_data['deletable_files'] = []
        message_lines = [f"*Pilih file yang ingin dihapus dari {category.upper()}:*\n"]
        sorted_files = sorted(files, key=lambda x: x['DateCreated'])
        for i, file_info in enumerate(sorted_files, 1):
            try:
                file_name = file_info['ObjectName']; parts = file_name.replace(".jpg", "").split("_")
                keterangan = " ".join(parts[1:-1]).capitalize(); nilai = parts[-1]
                message_lines.append(f"`{i}`. {keterangan} - Rp {int(nilai):,}")
                context.user_data['deletable_files'].append(file_name)
            except (ValueError, IndexError): continue
        message_lines.append("\nKetik nomor file yang ingin Anda hapus, atau /batal.")
        await query.edit_message_text("\n".join(message_lines), parse_mode='Markdown')
        return self.CHOOSE_DELETE_FILE

    async def choose_delete_file(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        try:
            choice = int(update.message.text); deletable_files = context.user_data.get('deletable_files', [])
            category = context.user_data.get('delete_category')
            if not category or not deletable_files:
                await update.message.reply_text("Sesi telah berakhir. Mulai lagi dengan /hapus."); context.user_data.clear(); return ConversationHandler.END
            if 0 < choice <= len(deletable_files):
                file_to_delete = deletable_files[choice - 1]; period, _, _ = self._get_current_period()
                remote_path = f"{category}/{period}/{file_to_delete}"
                await update.message.reply_text(f"⏳ Menghapus file `{file_to_delete}`...", parse_mode='Markdown')
                success, message = self.storage.delete_file(remote_path)
                await update.message.reply_text(message)
                context.user_data.clear(); return ConversationHandler.END
            else:
                await update.message.reply_text("Nomor tidak valid. Silakan pilih nomor dari daftar."); return self.CHOOSE_DELETE_FILE
        except ValueError:
            await update.message.reply_text("Input tidak valid. Harap masukkan nomor saja."); return self.CHOOSE_DELETE_FILE
        except Exception as e:
            logger.error(f"Error during file deletion choice: {e}"); await update.message.reply_text("Terjadi kesalahan saat proses penghapusan.")
            context.user_data.clear(); return ConversationHandler.END

    # --- Alur Percakapan Edit Data ---
    async def start_edit_flow(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        buttons = [[InlineKeyboardButton(c.upper(), callback_data=f"editcat_{c}")] for c in self.commands]
        await update.message.reply_text("Pilih kategori data yang ingin diedit:", reply_markup=InlineKeyboardMarkup(buttons))
        return self.CHOOSE_EDIT_CATEGORY

    async def choose_edit_category(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query; await query.answer(); category = query.data.split('_')[1]
        context.user_data['edit_category'] = category; period, _, _ = self._get_current_period()
        remote_path = f"{category}/{period}/"; files = self.storage.list_files(remote_path)
        if not files:
            await query.edit_message_text(f"Tidak ada data untuk diedit di *{category.upper()}*.", parse_mode='Markdown'); context.user_data.clear(); return ConversationHandler.END
        context.user_data['editable_files'] = sorted(files, key=lambda x: x['DateCreated'])
        message = [f"*Pilih data dari {category.upper()} yang ingin diedit:*\n"]
        for i, f in enumerate(context.user_data['editable_files'], 1):
            try:
                parts=f['ObjectName'].replace('.jpg','').split('_'); ket= " ".join(parts[1:-1]).capitalize(); nilai=parts[-1]
                message.append(f"`{i}`. {ket} - Rp {int(nilai):,}")
            except (ValueError, IndexError): continue
        message.append("\nKetik nomor data yang ingin diedit, atau /batal.")
        await query.edit_message_text("\n".join(message), parse_mode='Markdown'); return self.CHOOSE_EDIT_FILE

    async def choose_edit_file(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        try: choice = int(update.message.text)
        except ValueError: await update.message.reply_text("Harap masukkan nomor."); return self.CHOOSE_EDIT_FILE
        if not 0 < choice <= len(context.user_data['editable_files']):
            await update.message.reply_text("Nomor tidak valid."); return self.CHOOSE_EDIT_FILE
        context.user_data['edit_file_index'] = choice - 1
        buttons = [[InlineKeyboardButton("Keterangan", callback_data="editfield_ket"), InlineKeyboardButton("Biaya", callback_data="editfield_bia")]]
        await update.message.reply_text("Apa yang ingin Anda ubah?", reply_markup=InlineKeyboardMarkup(buttons)); return self.CHOOSE_EDIT_FIELD

    async def choose_edit_field(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query; await query.answer(); field = query.data.split('_')[1]
        context.user_data['edit_field'] = field
        await query.edit_message_text(f"Masukkan {field} baru:"); return self.GET_NEW_VALUE

    async def get_new_value_and_save(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        new_value = update.message.text; field_to_edit = context.user_data['edit_field']
        if field_to_edit == 'bia' and not new_value.isdigit():
            await update.message.reply_text("Biaya harus angka. Masukkan lagi."); return self.GET_NEW_VALUE
        await update.message.reply_text("⏳ Menyimpan perubahan...")
        try:
            category = context.user_data['edit_category']; period, _, _ = self._get_current_period()
            original_file = context.user_data['editable_files'][context.user_data['edit_file_index']]
            original_filename = original_file['ObjectName']
            parts = original_filename.replace(".jpg","").split("_")
            timestamp, original_keterangan, original_biaya = parts[0], "_".join(parts[1:-1]), parts[-1]
            new_keterangan = re.sub(r'[^\w\-_\.]', '_', new_value) if field_to_edit == 'ket' else original_keterangan
            new_biaya = new_value if field_to_edit == 'bia' else original_biaya
            new_filename = f"{timestamp}_{new_keterangan}_{new_biaya}.jpg"
            if new_filename == original_filename:
                await update.message.reply_text("Tidak ada perubahan. Proses dihentikan."); context.user_data.clear(); return ConversationHandler.END
            with tempfile.TemporaryDirectory() as temp_dir:
                local_path = os.path.join(temp_dir, original_filename)
                original_remote_path = f"{category}/{period}/{original_filename}"
                if self.storage.download_file(original_remote_path, local_path):
                    success, msg = self.storage.delete_file(original_remote_path)
                    if success:
                        new_remote_path = f"{category}/{period}/{new_filename}"
                        if self.storage.upload_file(local_path, new_remote_path):
                            await update.message.reply_text("✅ Perubahan berhasil disimpan.")
                        else: await update.message.reply_text("❌ Gagal mengunggah file baru.")
                    else: await update.message.reply_text(f"❌ Gagal menghapus file lama: {msg}")
                else: await update.message.reply_text("❌ Gagal mengunduh file untuk diedit.")
        except Exception as e:
            logger.error(f"Kesalahan saat edit data: {e}"); await update.message.reply_text("❌ Terjadi kesalahan internal.")
        finally:
            context.user_data.clear(); return ConversationHandler.END
    
    async def cancel(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        if context.user_data: context.user_data.clear()
        if update.callback_query:
            await update.callback_query.answer(); await update.callback_query.edit_message_text("Proses dibatalkan.")
        else:
            await update.message.reply_text("Proses dibatalkan.")
        return ConversationHandler.END

    async def list_data(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        period, year, month = self._get_current_period(); month_name = datetime(year, month, 1).strftime("%B")
        messages = [f"📋 *Data Rembesan Cloud - {month_name} {year}:*\n"]; total_data, total_biaya = 0, 0
        for cmd in self.commands:
            remote_path = f"{cmd}/{period}/"; files = self.storage.list_files(remote_path)
            if files:
                cmd_data, cmd_total = [], 0
                for file_info in sorted(files, key=lambda x: x['DateCreated']):
                    try:
                        file_name = file_info['ObjectName']; parts = file_name.replace(".jpg", "").split("_")
                        nilai, keterangan = parts[-1], " ".join(parts[1:-1]).capitalize(); nilai_int = int(nilai)
                        cmd_data.append(f"- {keterangan}: Rp {nilai_int:,}"); total_data += 1
                        total_biaya += nilai_int; cmd_total += nilai_int
                    except (ValueError, IndexError): continue
                if cmd_data:
                    messages.append(f"*{cmd.upper()}* (Total: Rp {cmd_total:,}):"); messages.extend(cmd_data); messages.append("")
        if total_data > 0:
            messages.append(f"*Total Transaksi: {total_data}*"); messages.append(f"*Total Biaya: Rp {total_biaya:,}*")
        else: messages.append("Belum ada data di cloud untuk periode ini.")
        await update.message.reply_text("\n".join(messages), parse_mode='Markdown')

    async def export_data(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        message = await update.message.reply_text("⏳ Memulai proses ekspor dari cloud...")
        temp_dir = tempfile.mkdtemp()
        try:
            period, year, month = self._get_current_period(); month_name = datetime(year, month, 1).strftime("%B")
            document = Document(); document.add_heading(f'Laporan Rembesan - {month_name} {year}', level=1)
            document.add_paragraph(f"Dibuat pada: {datetime.now().strftime('%d %B %Y, %H:%M:%S')}")
            grand_total, has_any_data = 0, False
            await context.bot.edit_message_text(chat_id=update.effective_chat.id, message_id=message.message_id, text="🔍 Mengambil data kategori...")
            for cmd in self.commands:
                remote_path = f"{cmd}/{period}/"; files = self.storage.list_files(remote_path)
                if not files: continue
                has_any_data, category_total = True, 0
                await context.bot.edit_message_text(chat_id=update.effective_chat.id, message_id=message.message_id, text=f"📄 Memproses kategori: {cmd.upper()}...")
                document.add_heading(cmd.upper(), level=2)
                table = document.add_table(rows=1, cols=4); table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'No.'; hdr_cells[1].text = 'Tanggal & Waktu'; hdr_cells[2].text = 'Keterangan & Bukti'; hdr_cells[3].text = 'Biaya (Rp)'
                sorted_files = sorted(files, key=lambda x: x['DateCreated'])
                for item_number, file_info in enumerate(sorted_files, 1):
                    try:
                        file_name = file_info['ObjectName']; parts = file_name.replace(".jpg", "").split("_")
                        timestamp_str, nilai_str, keterangan = parts[0], parts[-1], " ".join(parts[1:-1]).capitalize()
                        nilai_int = int(nilai_str); category_total += nilai_int
                        try: tanggal_formatted = datetime.strptime(timestamp_str, '%Y%m%d_%H%M%S').strftime('%d %b %Y, %H:%M')
                        except ValueError: tanggal_formatted = timestamp_str
                        row_cells = table.add_row().cells; row_cells[0].text = str(item_number); row_cells[1].text = tanggal_formatted
                        p = row_cells[2].paragraphs[0]; p.add_run(keterangan).bold = True
                        local_image_path = os.path.join(temp_dir, file_name)
                        if self.storage.download_file(f"{remote_path}{file_name}", local_image_path):
                            p.add_run('\n').add_picture(local_image_path, width=Inches(1.5))
                        row_cells[3].text = f"{nilai_int:,}"
                    except Exception as e: logger.error(f"Gagal memproses file {file_name} untuk ekspor: {e}")
                total_row = table.add_row(); merged_cell = total_row.cells[0].merge(total_row.cells[2])
                merged_cell.text = f"Total {cmd.upper()}"; merged_cell.paragraphs[0].runs[0].bold = True
                total_value_cell = total_row.cells[3]; total_value_cell.text = f"Rp {category_total:,}"; total_value_cell.paragraphs[0].runs[0].bold = True
                grand_total += category_total
            if not has_any_data:
                await context.bot.edit_message_text(chat_id=update.effective_chat.id, message_id=message.message_id, text="ℹ️ Tidak ada data untuk diekspor.")
                return
            document.add_heading('Ringkasan Total', level=2)
            document.add_paragraph().add_run(f"Grand Total: Rp {grand_total:,}").bold = True
            export_path = os.path.join(temp_dir, f'Laporan_Rembesan_{year}_{month:02d}.docx')
            document.save(export_path)
            await context.bot.edit_message_text(chat_id=update.effective_chat.id, message_id=message.message_id, text="✅ Ekspor berhasil! Mengirim dokumen...")
            await context.bot.send_document(chat_id=update.effective_chat.id, document=open(export_path, 'rb'))
            await context.bot.delete_message(chat_id=update.effective_chat.id, message_id=message.message_id)
        except Exception as e:
            logger.error(f"Kesalahan fatal saat ekspor data: {e}"); await context.bot.edit_message_text(chat_id=update.effective_chat.id, message_id=message.message_id, text=f"❌ Terjadi kesalahan: {e}")
        finally:
            shutil.rmtree(temp_dir)

    # --- Otomatisasi & Menjalankan Bot ---
    async def post_init(self, application: Application):
        """Menjalankan scheduler setelah aplikasi bot siap."""
        scheduler = AsyncIOScheduler(timezone="Asia/Jakarta")
        scheduler.add_job(self.send_reminder, 'cron', day=22, hour=9)
        scheduler.start()
        logger.info("Scheduler untuk pengingat otomatis telah dimulai.")

    async def send_reminder(self):
        period, _, _ = self._get_current_period()
        message = (
            f"🔔 *PENGINGAT REMBESAN* 🔔\n\n"
            f"Batas akhir pengajuan rembesan untuk periode *{period}* adalah tanggal 24 bulan ini.\n\n"
            "Mohon segera unggah semua bukti pembayaran Anda. Terima kasih!"
        )
        if self.config.GROUP_CHAT_ID:
            try:
                await self.application.bot.send_message(chat_id=self.config.GROUP_CHAT_ID, text=message, parse_mode='Markdown')
                logger.info("Pesan pengingat berhasil dikirim.")
            except Exception as e:
                logger.error(f"Gagal mengirim pesan pengingat: {e}")

    def run(self):
        if not all([self.config.TOKEN, self.config.BUNNY_STORAGE_ZONE_NAME, self.config.BUNNY_ACCESS_KEY]):
            logger.critical("TOKEN atau kredensial BUNNY tidak lengkap!"); return
        self.application = ApplicationBuilder().token(self.config.TOKEN).post_init(self.post_init).build()

        add_conv = ConversationHandler(
            entry_points=[CommandHandler(cmd, self.start_reimbursement_flow) for cmd in self.commands],
            states={
                self.GET_PHOTO: [MessageHandler(filters.PHOTO, self.get_photo)],
                self.GET_KETERANGAN: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.get_keterangan)],
                self.GET_BIAYA: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.get_biaya_and_save)],
                self.ASK_CONTINUE: [CallbackQueryHandler(self.ask_continue, pattern="^continue_")],
            }, fallbacks=[CommandHandler("batal", self.cancel)], per_message=True
        )
        delete_conv = ConversationHandler(
            entry_points=[CommandHandler("hapus", self.start_delete_flow)],
            states={
                self.CHOOSE_DELETE_CATEGORY: [CallbackQueryHandler(self.choose_delete_category, pattern="^delcat_")],
                self.CHOOSE_DELETE_FILE: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.choose_delete_file)],
            }, fallbacks=[CommandHandler("batal", self.cancel)], per_message=True
        )
        edit_conv = ConversationHandler(
            entry_points=[CommandHandler("edit", self.start_edit_flow)],
            states={
                self.CHOOSE_EDIT_CATEGORY: [CallbackQueryHandler(self.choose_edit_category, pattern="^editcat_")],
                self.CHOOSE_EDIT_FILE: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.choose_edit_file)],
                self.CHOOSE_EDIT_FIELD: [CallbackQueryHandler(self.choose_edit_field, pattern="^editfield_")],
                self.GET_NEW_VALUE: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.get_new_value_and_save)],
            }, fallbacks=[CommandHandler("batal", self.cancel)], per_message=True
        )

        self.application.add_handler(add_conv); self.application.add_handler(delete_conv); self.application.add_handler(edit_conv)
        self.application.add_handler(CommandHandler("start", self.start_command)); self.application.add_handler(CommandHandler("status", self.status_command))
        self.application.add_handler(CommandHandler("summary", self.summary_command)); self.application.add_handler(CommandHandler("list", self.list_data))
        self.application.add_handler(CommandHandler("export", self.export_data)); self.application.add_handler(CommandHandler("tambah_kategori", self.add_category_command))
        self.application.add_handler(CommandHandler("hapus_kategori", self.remove_category_command))
        
        logger.info("Bot v3.1 (Full Feature) dimulai...")
        self.application.run_polling()

if __name__ == "__main__":
    bot = RembesBot()
    bot.run()
